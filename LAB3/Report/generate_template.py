#!/usr/bin/env python3
"""
Script para gerar template Word no formato CBA/SBAI
Autor: Claude Code
Data: 2025
"""

import sys
import subprocess

# Verificar e instalar python-docx se necessário
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[INFO] Instalando biblioteca python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

def set_margins(section, top=2.5, bottom=2.5, left=2.5, right=2.5):
    """Define margens do documento em centímetros"""
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)

def create_cba_sbai_template():
    """Cria template Word no formato CBA/SBAI"""

    print("[INFO] Criando template CBA/SBAI...")

    # Criar documento
    doc = Document()

    # Configurar seção (página A4)
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4 altura
    section.page_width = Cm(21.0)   # A4 largura
    set_margins(section, 2.5, 2.5, 2.5, 2.5)

    # ========================================
    # CONFIGURAR ESTILOS
    # ========================================

    styles = doc.styles

    # --- Estilo: Normal (corpo do texto) ---
    try:
        style_normal = styles['Normal']
    except KeyError:
        style_normal = styles.add_style('Normal', 1)  # 1 = PARAGRAPH

    font_normal = style_normal.font
    font_normal.name = 'Times New Roman'
    font_normal.size = Pt(12)

    paragraph_format = style_normal.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5 linhas
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(6)
    paragraph_format.first_line_indent = Cm(0)

    # --- Estilo: Title (título principal) ---
    try:
        style_title = styles['Title']
    except KeyError:
        style_title = styles.add_style('Title', 1)

    font_title = style_title.font
    font_title.name = 'Times New Roman'
    font_title.size = Pt(14)
    font_title.bold = True
    font_title.color.rgb = RGBColor(0, 0, 0)

    paragraph_format_title = style_title.paragraph_format
    paragraph_format_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format_title.space_before = Pt(0)
    paragraph_format_title.space_after = Pt(12)

    # --- Estilo: Heading 1 (Seção principal: 1. INTRODUÇÃO) ---
    try:
        style_h1 = styles['Heading 1']
    except KeyError:
        style_h1 = styles.add_style('Heading 1', 1)

    font_h1 = style_h1.font
    font_h1.name = 'Times New Roman'
    font_h1.size = Pt(12)
    font_h1.bold = True
    font_h1.all_caps = True
    font_h1.color.rgb = RGBColor(0, 0, 0)

    paragraph_format_h1 = style_h1.paragraph_format
    paragraph_format_h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_format_h1.space_before = Pt(18)
    paragraph_format_h1.space_after = Pt(6)
    paragraph_format_h1.keep_with_next = True

    # --- Estilo: Heading 2 (Subseção: 2.1. Modelo FOPDT) ---
    try:
        style_h2 = styles['Heading 2']
    except KeyError:
        style_h2 = styles.add_style('Heading 2', 1)

    font_h2 = style_h2.font
    font_h2.name = 'Times New Roman'
    font_h2.size = Pt(12)
    font_h2.bold = True
    font_h2.color.rgb = RGBColor(0, 0, 0)

    paragraph_format_h2 = style_h2.paragraph_format
    paragraph_format_h2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_format_h2.space_before = Pt(12)
    paragraph_format_h2.space_after = Pt(6)
    paragraph_format_h2.keep_with_next = True

    # --- Estilo: Heading 3 (Sub-subseção: 2.1.1.) ---
    try:
        style_h3 = styles['Heading 3']
    except KeyError:
        style_h3 = styles.add_style('Heading 3', 1)

    font_h3 = style_h3.font
    font_h3.name = 'Times New Roman'
    font_h3.size = Pt(12)
    font_h3.bold = True
    font_h3.italic = True
    font_h3.color.rgb = RGBColor(0, 0, 0)

    paragraph_format_h3 = style_h3.paragraph_format
    paragraph_format_h3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_format_h3.space_before = Pt(6)
    paragraph_format_h3.space_after = Pt(6)

    # --- Estilo: Caption (legenda de figura/tabela) ---
    try:
        style_caption = styles['Caption']
    except KeyError:
        style_caption = styles.add_style('Caption', 1)

    font_caption = style_caption.font
    font_caption.name = 'Times New Roman'
    font_caption.size = Pt(11)
    font_caption.color.rgb = RGBColor(0, 0, 0)

    paragraph_format_caption = style_caption.paragraph_format
    paragraph_format_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format_caption.space_before = Pt(6)
    paragraph_format_caption.space_after = Pt(12)

    # --- Estilo: Subtitle (subtítulo, autores) ---
    try:
        style_subtitle = styles['Subtitle']
    except KeyError:
        style_subtitle = styles.add_style('Subtitle', 1)

    font_subtitle = style_subtitle.font
    font_subtitle.name = 'Times New Roman'
    font_subtitle.size = Pt(12)
    font_subtitle.bold = True
    font_subtitle.color.rgb = RGBColor(0, 0, 0)

    paragraph_format_subtitle = style_subtitle.paragraph_format
    paragraph_format_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format_subtitle.space_before = Pt(0)
    paragraph_format_subtitle.space_after = Pt(6)

    # --- Estilo Customizado: Affiliation (afiliação em itálico) ---
    try:
        style_affiliation = styles.add_style('Affiliation', 1)
    except:
        style_affiliation = styles['Affiliation']

    font_affiliation = style_affiliation.font
    font_affiliation.name = 'Times New Roman'
    font_affiliation.size = Pt(11)
    font_affiliation.italic = True
    font_affiliation.color.rgb = RGBColor(0, 0, 0)

    paragraph_format_affiliation = style_affiliation.paragraph_format
    paragraph_format_affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format_affiliation.space_before = Pt(0)
    paragraph_format_affiliation.space_after = Pt(3)

    # --- Estilo: Intense Quote (RESUMO destacado) ---
    try:
        style_abstract = styles['Intense Quote']
    except KeyError:
        style_abstract = styles.add_style('Intense Quote', 1)

    font_abstract = style_abstract.font
    font_abstract.name = 'Times New Roman'
    font_abstract.size = Pt(12)

    paragraph_format_abstract = style_abstract.paragraph_format
    paragraph_format_abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format_abstract.space_before = Pt(6)
    paragraph_format_abstract.space_after = Pt(12)
    paragraph_format_abstract.left_indent = Cm(0)
    paragraph_format_abstract.right_indent = Cm(0)

    # ========================================
    # ADICIONAR CONTEÚDO DE EXEMPLO
    # ========================================

    # Título
    p_title = doc.add_paragraph('Título do Artigo Científico no Formato CBA/SBAI', style='Title')

    # Autores
    p_authors = doc.add_paragraph('Nome do Autor 1, Nome do Autor 2', style='Subtitle')

    # Afiliação
    p_aff1 = doc.add_paragraph(
        'Programa de Pós-Graduação em Engenharia (PPGEP)',
        style='Affiliation'
    )
    p_aff2 = doc.add_paragraph(
        'Universidade Federal do Brasil (UFB)',
        style='Affiliation'
    )
    p_aff3 = doc.add_paragraph(
        'Rua Exemplo, 123, CEP 12345-678, Cidade, Estado, Brasil',
        style='Affiliation'
    )

    # Email
    p_email = doc.add_paragraph()
    p_email.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_email = p_email.add_run('E-mail: autor1@universidade.br, autor2@universidade.br')
    run_email.font.name = 'Times New Roman'
    run_email.font.size = Pt(10)

    # Linha horizontal
    doc.add_paragraph('_' * 80)

    # RESUMO
    doc.add_heading('RESUMO', level=1)
    doc.add_paragraph(
        'Este é um exemplo de resumo. O resumo deve conter entre 150-250 palavras, '
        'apresentando de forma concisa os objetivos, metodologia, resultados principais '
        'e conclusões do trabalho. Deve ser escrito em parágrafo único, sem citações, '
        'em português (para artigos em português) ou inglês (para artigos em inglês).',
        style='Normal'
    )

    p_keywords = doc.add_paragraph()
    p_keywords.add_run('Palavras-chave: ').bold = True
    p_keywords.add_run('Palavra1, Palavra2, Palavra3, Palavra4.')
    p_keywords.style = 'Normal'

    # Seções
    doc.add_heading('1. INTRODUÇÃO', level=1)
    doc.add_paragraph(
        'A introdução contextualiza o tema do trabalho, apresenta a motivação, '
        'os objetivos e a estrutura do artigo. Deve ser clara e objetiva.',
        style='Normal'
    )

    doc.add_heading('2. FUNDAMENTAÇÃO TEÓRICA', level=1)
    doc.add_paragraph(
        'Esta seção apresenta a base teórica necessária para compreensão do trabalho.',
        style='Normal'
    )

    doc.add_heading('2.1. Subseção de Teoria', level=2)
    doc.add_paragraph(
        'As subseções organizam o conteúdo de forma lógica e hierárquica.',
        style='Normal'
    )

    doc.add_heading('3. METODOLOGIA', level=1)
    doc.add_paragraph(
        'Descreve os métodos, procedimentos e técnicas utilizados no trabalho.',
        style='Normal'
    )

    doc.add_heading('4. RESULTADOS', level=1)
    doc.add_paragraph(
        'Apresenta os resultados obtidos, com tabelas, figuras e análises.',
        style='Normal'
    )

    # Exemplo de tabela
    doc.add_paragraph('Tabela 1: Exemplo de tabela formatada', style='Caption')
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Light Grid Accent 1'

    # Cabeçalho da tabela
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Parâmetro'
    header_cells[1].text = 'Valor 1'
    header_cells[2].text = 'Valor 2'
    header_cells[3].text = 'Valor 3'

    # Dados
    row1_cells = table.rows[1].cells
    row1_cells[0].text = 'K'
    row1_cells[1].text = '1.00'
    row1_cells[2].text = '1.50'
    row1_cells[3].text = '2.00'

    row2_cells = table.rows[2].cells
    row2_cells[0].text = 'T'
    row2_cells[1].text = '0.50'
    row2_cells[2].text = '0.75'
    row2_cells[3].text = '1.00'

    doc.add_paragraph()  # Espaço

    # Exemplo de figura
    doc.add_paragraph(
        'Figura 1: Descrição da figura. Figuras devem ser inseridas com boa resolução '
        'e legendas descritivas.',
        style='Caption'
    )

    doc.add_paragraph()  # Espaço

    doc.add_heading('5. DISCUSSÃO', level=1)
    doc.add_paragraph(
        'Análise e interpretação dos resultados obtidos.',
        style='Normal'
    )

    doc.add_heading('6. CONCLUSÃO', level=1)
    doc.add_paragraph(
        'Síntese dos principais resultados, contribuições do trabalho e sugestões '
        'para trabalhos futuros.',
        style='Normal'
    )

    doc.add_heading('REFERÊNCIAS', level=1)
    doc.add_paragraph(
        '[1] SOBRENOME, Iniciais. Título do Livro. Edição. Cidade: Editora, Ano.',
        style='Normal'
    )
    doc.add_paragraph(
        '[2] SOBRENOME, Iniciais. Título do artigo. Nome da Revista, v. X, n. Y, '
        'p. ZZ-WW, ano.',
        style='Normal'
    )

    # Salvar template
    output_file = 'template_cba_sbai.docx'
    doc.save(output_file)

    print(f"\n[SUCESSO] Template criado: {output_file}")
    print(f"\nConfiguracoes do template:")
    print(f"   - Margens: 2.5 cm (todas)")
    print(f"   - Fonte: Times New Roman 12pt")
    print(f"   - Espacamento: 1.5 linhas")
    print(f"   - Tamanho: A4")
    print(f"\nProximo passo - Use o comando:")
    print(f"   pandoc LAB3_Relatorio_Tecnico_Victor_Ramon.md -o LAB3_Relatorio_Tecnico_Victor_Ramon.docx --reference-doc={output_file} --number-sections")

if __name__ == "__main__":
    create_cba_sbai_template()
